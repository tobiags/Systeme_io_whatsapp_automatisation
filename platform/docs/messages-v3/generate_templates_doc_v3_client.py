from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Deprecated client-facing generator.
# Kept for archive/history only.
# Active client validation base is now:
# platform/docs/messages-v2.1/messages-challenge-amazon-fba-v2.1-validation-client.docx


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "messages-challenge-amazon-fba-v3-validation-client.docx"
OUTPUT_MD = ROOT / "messages-challenge-amazon-fba-v3-validation-client.md"

ACCENT = RGBColor(0x1D, 0x4E, 0x89)
ACCENT_DARK = RGBColor(0x15, 0x3B, 0x6D)
TEXT = RGBColor(0x1F, 0x29, 0x37)
TABLE_HEAD = "D9E8FB"
INFO = "EAF1FB"
SUCCESS = "EAF6EE"
BORDER = "CBD5E1"

SESSION_1 = "La méthode FBA de A à Z"
SESSION_2 = "Construire Ton Business Amazon Pas à Pas"
SESSION_3 = "Le Secret des Vendeurs à Succès sur Amazon"


@dataclass
class Variable:
    key: str
    description: str
    sample: str


@dataclass
class Template:
    key: str
    title: str
    phase: str
    timing: str
    audience: str
    variables: list[Variable]
    body: list[str]
    validation_points: list[str] = field(default_factory=list)


INTRO_POINTS = [
    "Cette version V3 intègre les derniers retours du client sur la cadence des rappels, les formulations à corriger et les messages de fin de séquence.",
    "Le ton retenu reste celui du formateur qui écrit directement : humain, simple, direct et jamais trop marketé.",
    "Les formulations trop marquées par le genre ont été retirées au profit de formulations neutres comme 'ta présence' ou 'je ne t'ai pas vu'.",
    "Ce document sert uniquement à valider les messages avant soumission finale dans Wati.",
]


TEMPLATES: list[Template] = [
    Template(
        key="welcome",
        title="Bienvenue",
        phase="Pré-challenge",
        timing="Immédiatement après l'inscription",
        audience="Tous les inscrits",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "Je suis ravi de t'accueillir pour le Challenge Amazon FBA.",
            "",
            "Pendant les prochains jours, je vais t'envoyer ici l'essentiel pour te préparer aux 3 directs.",
            "",
            "Dis-moi : as-tu déjà entendu parler de la vente sur Amazon ?",
        ],
    ),
    Template(
        key="countdown_j6",
        title="Question sur la situation actuelle",
        phase="Pré-challenge",
        timing="J-6",
        audience="Tous les inscrits",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "Avant qu'on démarre, j'ai une question simple pour toi.",
            "",
            "Aujourd'hui, ton plus gros frein pour te lancer sur Amazon FBA, c'est quoi ?",
            "",
            "Le temps, le budget, le choix du produit, ou autre chose ? Réponds-moi ici.",
        ],
    ),
    Template(
        key="countdown_j5",
        title="Curiosité",
        phase="Pré-challenge",
        timing="J-5",
        audience="Tous les inscrits",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "Petit point avant la suite.",
            "",
            "Beaucoup de vendeurs Amazon ne touchent jamais leurs produits. Amazon stocke, expédie et gère une grosse partie de l'opération.",
            "",
            "C'est justement ce qu'on va remettre à plat ensemble.",
            "",
            "De ton côté, qu'est-ce qui t'intrigue le plus dans ce modèle ?",
        ],
    ),
    Template(
        key="countdown_j4",
        title="Communauté et engagement",
        phase="Pré-challenge",
        timing="J-4",
        audience="Tous les inscrits",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "On a déjà un très bon groupe pour cette édition, avec des profils très différents.",
            "",
            "Certains partent de zéro, d'autres ont déjà essayé de vendre en ligne, mais tous viennent avec la même envie : comprendre comment lancer quelque chose de sérieux.",
            "",
            "Tu verras, on va avancer simplement et concrètement.",
        ],
    ),
    Template(
        key="countdown_j3",
        title="Programme des 3 directs",
        phase="Pré-challenge",
        timing="J-3",
        audience="Tous les inscrits",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "Voici ce qu'on va voir ensemble pendant les 3 directs :",
            "",
            f"1. {SESSION_1}",
            f"2. {SESSION_2}",
            f"3. {SESSION_3}",
            "",
            "Pense à t'inscrire en avance s'il te plaît.",
            "",
            "Seras-tu parmi nous en direct ?",
        ],
    ),
    Template(
        key="countdown_j2",
        title="Informations pratiques",
        phase="Pré-challenge",
        timing="J-2",
        audience="Tous les inscrits",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "Dans 2 jours, on commence.",
            "",
            "Quelques points simples :",
            "- le live se fait sur StreamYard",
            "- tu recevras ton lien ici avant chaque session",
            "- pas besoin de créer de compte",
            "- si tu arrives 5 minutes en avance, c'est parfait",
            "",
            "Si tu as une question pratique avant le début, réponds-moi ici.",
        ],
    ),
    Template(
        key="countdown_j1",
        title="La veille du premier live",
        phase="Pré-challenge",
        timing="J-1",
        audience="Tous les inscrits",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Heure de la session du Jour 1", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On démarre demain à {{2}}.",
            "",
            "Je t'enverrai le lien ici avant le live.",
            "",
            "Prends juste le créneau au calme si tu peux, parce que la première session pose toute la base pour la suite.",
            "",
            "À demain.",
        ],
    ),
    Template(
        key="live_day1_h2",
        title="Jour 1 - rappel principal",
        phase="Jour 1",
        timing="2 heures avant le live",
        audience="Tous les inscrits actifs",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 1", "https://streamyard.com/watch/exemple-j1"),
            Variable("{{3}}", "Heure du live Jour 1", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            f"On se retrouve ce soir à {{{{3}}}} pour : {SESSION_1}",
            "",
            "Voici ton lien pour nous rejoindre :",
            "{{2}}",
            "",
            "Essaie d'être là dès le début, cette première session pose toute la base.",
        ],
    ),
    Template(
        key="live_day1_h10",
        title="Jour 1 - rappel 10 minutes avant",
        phase="Jour 1",
        timing="10 minutes avant le live",
        audience="Tous les inscrits actifs",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 1", "https://streamyard.com/watch/exemple-j1"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On démarre dans 10 minutes.",
            "",
            "Voici ton lien :",
            "{{2}}",
            "",
            "À tout de suite.",
        ],
    ),
    Template(
        key="live_day1_hplus5",
        title="Jour 1 - relance 5 minutes après le début",
        phase="Jour 1",
        timing="5 minutes après le début du live",
        audience="Ceux qui n'ont pas cliqué ou ouvert le message précédent",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 1", "https://streamyard.com/watch/exemple-j1"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On a commencé il y a 5 minutes.",
            "",
            "Si tu comptais nous rejoindre, tu peux encore entrer maintenant :",
            "{{2}}",
        ],
    ),
    Template(
        key="live_day2_attended_h2",
        title="Jour 2 - 2h avant pour les présents du Jour 1",
        phase="Jour 2",
        timing="2 heures avant le live",
        audience="Présents au Jour 1",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 2", "https://streamyard.com/watch/exemple-j2"),
            Variable("{{3}}", "Heure du live Jour 2", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Merci pour ta présence hier.",
            "",
            f"Ce soir à {{{{3}}}}, on continue avec : {SESSION_2}",
            "",
            "Ton lien :",
            "{{2}}",
            "",
            "Hier, quel point t'a le plus marqué ?",
        ],
    ),
    Template(
        key="live_day2_registered_absent_h2",
        title="Jour 2 - 2h avant pour les inscrits absents au Jour 1",
        phase="Jour 2",
        timing="2 heures avant le live",
        audience="Inscrits StreamYard Jour 1 mais absents",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 2", "https://streamyard.com/watch/exemple-j2"),
            Variable("{{3}}", "Heure du live Jour 2", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Je ne t'ai pas vu hier soir.",
            "",
            f"Ce soir à {{{{3}}}}, on continue avec : {SESSION_2}",
            "",
            "Voici ton lien :",
            "{{2}}",
            "",
            "Dis-moi simplement ce qui t'a empêché d'être là hier.",
        ],
    ),
    Template(
        key="live_day2_not_registered_h2",
        title="Jour 2 - 2h avant pour les non inscrits au Jour 1",
        phase="Jour 2",
        timing="2 heures avant le live",
        audience="Aucune interaction StreamYard Jour 1",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 2", "https://streamyard.com/watch/exemple-j2"),
            Variable("{{3}}", "Heure du live Jour 2", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Si tu n'as pas pu suivre le premier direct, tu peux encore prendre le train en marche.",
            "",
            f"Ce soir à {{{{3}}}} : {SESSION_2}",
            "",
            "Voici le lien :",
            "{{2}}",
        ],
    ),
    Template(
        key="live_day2_h10",
        title="Jour 2 - rappel 10 minutes avant",
        phase="Jour 2",
        timing="10 minutes avant le live",
        audience="Tous les contacts ciblés pour Jour 2",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 2", "https://streamyard.com/watch/exemple-j2"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On démarre dans 10 minutes.",
            "",
            "Voici ton lien :",
            "{{2}}",
            "",
            "À tout de suite.",
        ],
    ),
    Template(
        key="live_day2_hplus5",
        title="Jour 2 - relance 5 minutes après le début",
        phase="Jour 2",
        timing="5 minutes après le début du live",
        audience="Ceux qui n'ont pas cliqué ou ouvert le message précédent",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 2", "https://streamyard.com/watch/exemple-j2"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On a commencé il y a 5 minutes.",
            "",
            "Si tu voulais nous rejoindre, tu peux encore entrer maintenant :",
            "{{2}}",
        ],
    ),
    Template(
        key="live_day3_attended_h2",
        title="Jour 3 - 2h avant pour les présents du Jour 2",
        phase="Jour 3",
        timing="2 heures avant le live",
        audience="Présents au Jour 2",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 3", "https://streamyard.com/watch/exemple-j3"),
            Variable("{{3}}", "Heure du live Jour 3", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Merci pour ta présence hier.",
            "",
            f"Ce soir à {{{{3}}}} : {SESSION_3}",
            "",
            "Ton lien :",
            "{{2}}",
            "",
            "Hier, quelle idée t'a le plus parlé ?",
        ],
    ),
    Template(
        key="live_day3_registered_absent_h2",
        title="Jour 3 - 2h avant pour les inscrits absents au Jour 2",
        phase="Jour 3",
        timing="2 heures avant le live",
        audience="Inscrits StreamYard Jour 2 mais absents",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 3", "https://streamyard.com/watch/exemple-j3"),
            Variable("{{3}}", "Heure du live Jour 3", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Je ne t'ai pas vu hier.",
            "",
            f"Ce soir à {{{{3}}}} : {SESSION_3}",
            "",
            "Ton lien :",
            "{{2}}",
            "",
            "Essaie vraiment d'être là, c'est la session la plus importante.",
        ],
    ),
    Template(
        key="live_day3_not_registered_h2",
        title="Jour 3 - 2h avant pour les non inscrits aux jours précédents",
        phase="Jour 3",
        timing="2 heures avant le live",
        audience="Aucune interaction StreamYard Jour 1 ou Jour 2",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 3", "https://streamyard.com/watch/exemple-j3"),
            Variable("{{3}}", "Heure du live Jour 3", "21h00"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Même si tu n'as pas encore suivi les directs, tu peux encore nous rejoindre ce soir.",
            "",
            f"Dernière session à {{{{3}}}} : {SESSION_3}",
            "",
            "Voici le lien :",
            "{{2}}",
        ],
    ),
    Template(
        key="live_day3_h10",
        title="Jour 3 - rappel 10 minutes avant",
        phase="Jour 3",
        timing="10 minutes avant le live",
        audience="Tous les contacts ciblés pour Jour 3",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 3", "https://streamyard.com/watch/exemple-j3"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On démarre dans 10 minutes.",
            "",
            "Voici ton lien :",
            "{{2}}",
            "",
            "À tout de suite.",
        ],
    ),
    Template(
        key="live_day3_hplus5",
        title="Jour 3 - relance 5 minutes après le début",
        phase="Jour 3",
        timing="5 minutes après le début du live",
        audience="Ceux qui n'ont pas cliqué ou ouvert le message précédent",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien StreamYard Jour 3", "https://streamyard.com/watch/exemple-j3"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "On a commencé il y a 5 minutes.",
            "",
            "Si tu comptais nous rejoindre, tu peux encore entrer maintenant :",
            "{{2}}",
        ],
    ),
    Template(
        key="live_day3_offer_hplus2",
        title="Jour 3 - lien de paiement 2 heures après le début",
        phase="Jour 3",
        timing="2 heures après le début du live",
        audience="Inscrits StreamYard Jour 3 n'ayant pas encore acheté",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien de paiement ou de programme", "https://paiement.exemple.com/programme"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Comme promis, voici le lien dont j'ai parlé pendant le direct :",
            "{{2}}",
            "",
            "Prends-le tranquillement.",
            "",
            "Et si tu veux poser une vraie question avant de décider, réponds-moi ici.",
        ],
        validation_points=[
            "Ce message doit être coupé automatiquement pour les contacts qui ont déjà payé dans la journée.",
        ],
    ),
    Template(
        key="post_replay_attended",
        title="J+1 - replay pour les plus engagés",
        phase="Post-challenge",
        timing="J+1",
        audience="Présents à la session 3",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Replay Jour 1", "https://replay.exemple.com/j1"),
            Variable("{{3}}", "Replay Jour 2", "https://replay.exemple.com/j2"),
            Variable("{{4}}", "Replay Jour 3", "https://replay.exemple.com/j3"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Merci encore pour ta présence pendant ce challenge.",
            "",
            "Comme promis, voici les replays :",
            "Jour 1 : {{2}}",
            "Jour 2 : {{3}}",
            "Jour 3 : {{4}}",
            "",
            "Prends le temps de revoir les points clés.",
        ],
        validation_points=[
            "À remplacer par les vrais liens replay transmis par le client.",
        ],
    ),
    Template(
        key="post_replay_partial",
        title="J+1 - replay pour rattrapage",
        phase="Post-challenge",
        timing="J+1",
        audience="Présents partiels ou absents partiels",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Replay Jour 1", "https://replay.exemple.com/j1"),
            Variable("{{3}}", "Replay Jour 2", "https://replay.exemple.com/j2"),
            Variable("{{4}}", "Replay Jour 3", "https://replay.exemple.com/j3"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Comme promis, voici les replays pour rattraper ce que tu as manqué :",
            "Jour 1 : {{2}}",
            "Jour 2 : {{3}}",
            "Jour 3 : {{4}}",
            "",
            "Regarde-les à ton rythme, puis dis-moi quel point t'a le plus aidé.",
        ],
        validation_points=[
            "À remplacer par les vrais liens replay transmis par le client.",
        ],
    ),
    Template(
        key="post_replay_absent",
        title="J+1 - replay pour les absents",
        phase="Post-challenge",
        timing="J+1",
        audience="Aucun live suivi",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Replay Jour 1", "https://replay.exemple.com/j1"),
            Variable("{{3}}", "Replay Jour 2", "https://replay.exemple.com/j2"),
            Variable("{{4}}", "Replay Jour 3", "https://replay.exemple.com/j3"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Tu n'as pas pu suivre les directs, donc je te laisse quand même les replays ici :",
            "Jour 1 : {{2}}",
            "Jour 2 : {{3}}",
            "Jour 3 : {{4}}",
            "",
            "Regarde-les à ton rythme, puis dis-moi si tu veux qu'on fasse le point.",
        ],
        validation_points=[
            "Ce message remplace la logique vague du 'réponds INFO'.",
            "À remplacer par les vrais liens replay transmis par le client.",
        ],
    ),
    Template(
        key="post_testimonials",
        title="Partage de témoignages",
        phase="Post-challenge",
        timing="J+2 ou J+3",
        audience="Contacts encore actifs",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "Je te partage deux retours de personnes qui sont passées à l'action après le challenge :",
            "",
            "[TÉMOIGNAGE 1 COURT]",
            "",
            "[TÉMOIGNAGE 2 COURT]",
            "",
            "Si tu veux qu'on voie si c'est possible pour toi aussi, réponds-moi simplement ici.",
        ],
        validation_points=[
            "À remplacer par 2 témoignages courts validés par le client.",
        ],
    ),
    Template(
        key="post_inaction_reason",
        title="Comprendre le frein réel",
        phase="Post-challenge",
        timing="J+3 ou J+4",
        audience="Contacts qui n'ont pas encore avancé",
        variables=[Variable("{{1}}", "Prénom du contact", "Sonia")],
        body=[
            "Bonjour {{1}},",
            "",
            "J'aimerais comprendre quelque chose.",
            "",
            "Qu'est-ce qui t'empêche aujourd'hui de passer à l'action : le budget, le temps, la peur de te tromper, ou autre chose ?",
            "",
            "Réponds-moi franchement.",
        ],
    ),
    Template(
        key="post_closer_call",
        title="Prise de rendez-vous closer",
        phase="Post-challenge",
        timing="Dernière relance",
        audience="Contacts encore chauds",
        variables=[
            Variable("{{1}}", "Prénom du contact", "Sonia"),
            Variable("{{2}}", "Lien formulaire closer / OnceHub", "https://www.ecommercecentrale.com/formulaire-challenge"),
        ],
        body=[
            "Bonjour {{1}},",
            "",
            "Si tu veux qu'on regarde ta situation plus sérieusement, tu peux réserver un échange ici :",
            "{{2}}",
            "",
            "Tu pourras poser tes questions et voir si l'accompagnement est vraiment fait pour toi.",
        ],
    ),
]


def set_table_borders(table, color: str = BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = TEXT
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.12

    for level, size in [(1, 20), (2, 14), (3, 11)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(8 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(4)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CHALLENGE AMAZON FBA")
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Séquence de messages WhatsApp - Version 3")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Document de validation client").font.size = Pt(12)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Projet", "Challenge Amazon FBA"),
        ("Version", "V3"),
        ("Voix retenue", "Le formateur écrit directement"),
        ("Statut", "À valider avant soumission Wati"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], TABLE_HEAD)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(table)
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    doc.add_paragraph("1. Ce qui change dans cette version", style="Heading 1")
    for point in INTRO_POINTS:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.add_run("• " + point)

    doc.add_paragraph("2. Structure retenue", style="Heading 1")
    structure = doc.add_table(rows=1, cols=2)
    structure.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(["Phase", "Contenu retenu"]):
        structure.rows[0].cells[index].text = header
        shade_cell(structure.rows[0].cells[index], TABLE_HEAD)
        structure.rows[0].cells[index].paragraphs[0].runs[0].bold = True
    rows = [
        ("Pré-challenge", "7 messages de préparation"),
        ("Jour 1", "3 messages : 2h avant, 10 min avant, 5 min après"),
        ("Jour 2", "5 messages : 3 branches à 2h + 10 min avant + 5 min après"),
        ("Jour 3", "6 messages : 3 branches à 2h + 10 min avant + 5 min après + lien de paiement"),
        ("Post-challenge", "6 messages : replays, témoignages, frein, prise de rendez-vous"),
    ]
    for left, right in rows:
        row = structure.add_row().cells
        row[0].text = left
        row[1].text = right
    set_table_borders(structure)

    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    note.rows[0].cells[0].text = (
        "Note : les messages de replay, les témoignages et certains liens restent à valider avec les éléments définitifs du client."
    )
    shade_cell(note.rows[0].cells[0], SUCCESS)
    set_table_borders(note)
    doc.add_page_break()


def templates_by_phase() -> dict[str, list[Template]]:
    phases: dict[str, list[Template]] = {}
    for template in TEMPLATES:
        phases.setdefault(template.phase, []).append(template)
    return phases


def add_template_section(doc: Document, template: Template, idx: int, total: int) -> None:
    doc.add_paragraph(f"Message {idx}/{total} - {template.title}", style="Heading 2")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Clé Wati", template.key),
        ("Phase", template.phase),
        ("Moment d'envoi", template.timing),
        ("Public concerné", template.audience),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], TABLE_HEAD)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(meta)

    doc.add_paragraph("Variables du message", style="Heading 3")
    variables = doc.add_table(rows=1, cols=3)
    variables.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(["Variable", "Usage", "Exemple"]):
        variables.rows[0].cells[index].text = header
        shade_cell(variables.rows[0].cells[index], TABLE_HEAD)
        variables.rows[0].cells[index].paragraphs[0].runs[0].bold = True
    for item in template.variables:
        row = variables.add_row().cells
        row[0].text = item.key
        row[1].text = item.description
        row[2].text = item.sample
    set_table_borders(variables)

    doc.add_paragraph("Texte du message WhatsApp", style="Heading 3")
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade_cell(box.rows[0].cells[0], INFO)
    set_table_borders(box)
    cell = box.rows[0].cells[0]
    for line in template.body:
        cell.add_paragraph(line)
    if cell.paragraphs and not cell.paragraphs[0].text.strip():
        cell._tc.remove(cell.paragraphs[0]._p)

    if template.validation_points:
        doc.add_paragraph("Point(s) à valider", style="Heading 3")
        for point in template.validation_points:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.add_run("• " + point)


def render_docx() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_intro(doc)

    ordered_phases = ["Pré-challenge", "Jour 1", "Jour 2", "Jour 3", "Post-challenge"]
    grouped = templates_by_phase()
    total = len(TEMPLATES)
    idx = 1
    for phase in ordered_phases:
        if phase not in grouped:
            continue
        doc.add_paragraph(phase, style="Heading 1")
        for template in grouped[phase]:
            add_template_section(doc, template, idx, total)
            idx += 1
    doc.save(OUTPUT_DOCX)


def render_markdown() -> None:
    lines: list[str] = [
        "# Challenge Amazon FBA",
        "",
        "## Séquence de messages WhatsApp - Version 3",
        "",
        "## Document de validation client",
        "",
        "### Ce que cette version change",
        "",
    ]
    lines.extend([f"- {point}" for point in INTRO_POINTS])
    lines.extend(
        [
            "",
            "### Structure retenue",
            "",
            "- Pré-challenge : 7 messages de préparation",
            "- Jour 1 : 3 messages (2h avant, 10 min avant, 5 min après)",
            "- Jour 2 : 5 messages (3 branches à 2h + 10 min avant + 5 min après)",
            "- Jour 3 : 6 messages (3 branches à 2h + 10 min avant + 5 min après + lien de paiement)",
            "- Post-challenge : 6 messages (replays, témoignages, frein, prise de rendez-vous)",
            "",
        ]
    )

    ordered_phases = ["Pré-challenge", "Jour 1", "Jour 2", "Jour 3", "Post-challenge"]
    grouped = templates_by_phase()
    total = len(TEMPLATES)
    idx = 1
    for phase in ordered_phases:
        if phase not in grouped:
            continue
        lines.append(f"## {phase}")
        lines.append("")
        for template in grouped[phase]:
            lines.append(f"### Message {idx}/{total} - {template.title}")
            lines.append("")
            lines.append(f"- Clé Wati : {template.key}")
            lines.append(f"- Moment d'envoi : {template.timing}")
            lines.append(f"- Public concerné : {template.audience}")
            lines.append("")
            lines.append("Variables :")
            for item in template.variables:
                lines.append(f"- {item.key} : {item.description} (exemple : {item.sample})")
            lines.append("")
            lines.append("```text")
            lines.extend(template.body)
            lines.append("```")
            if template.validation_points:
                lines.append("")
                lines.append("Point(s) à valider :")
                for point in template.validation_points:
                    lines.append(f"- {point}")
            lines.append("")
            idx += 1

    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    render_docx()
    render_markdown()
    print(f"Generated {OUTPUT_DOCX}")
    print(f"Generated {OUTPUT_MD}")


if __name__ == "__main__":
    main()
