from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "SEQUENCE-WHATSAPP-V3-WORKING.md"
OUTPUT = ROOT / "messages-challenge-amazon-fba-v3-working.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_LIGHT = "D9EAF7"
BORDER = "C9D3DD"


def set_table_borders(table, color: str = BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    for level, size in [(1, 20), (2, 15), (3, 12)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(10 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sequence WhatsApp V3")
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Base de travail apres retour client")
    run.font.name = "Aptos"
    run.font.size = Pt(15)
    run.font.bold = True

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Projet", "Challenge Amazon FBA"),
        ("Type", "Reprise V3 sequence et templates"),
        ("Statut", "Working draft"),
        ("Genere le", datetime.now().strftime("%d/%m/%Y a %H:%M")),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], ACCENT_LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(table)
    doc.add_page_break()


def parse_markdown(doc: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style="Heading 1")
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 2")
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 3")
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.add_run("• " + stripped[2:])
            continue
        if stripped[:2].isdigit() and stripped[1:3] == ". ":
            doc.add_paragraph(stripped)
            continue
        doc.add_paragraph(stripped)


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    main()
