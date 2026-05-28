from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = Path(__file__).with_name("MANUEL-UTILISATION-COMPLET-CHALLENGE-AMAZON-FBA.docx")
SCREENSHOTS = ROOT / "docs" / "screenshots"


COLORS = {
    "navy": RGBColor(11, 37, 69),
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "muted": RGBColor(89, 89, 89),
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "soft_green": "EAF7F0",
    "soft_yellow": "FFF7DB",
    "soft_red": "FDECEC",
    "border": "D9E2EC",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Manuel utilisateur - Challenge Amazon FBA | Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 8),
        ("Heading 2", 13, COLORS["blue"], 13, 6),
        ("Heading 3", 11.5, COLORS["dark_blue"], 9, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("MANUEL D'UTILISATION COMPLET")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = COLORS["navy"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plateforme WhatsApp - Challenge Amazon FBA")
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = COLORS["blue"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Guide client de A a Z : pages, liens, roles, procedures, verification et limites connues")
    r.font.size = Pt(11)
    r.font.color.rgb = COLORS["muted"]

    meta = [
        ("Version", "1.0 - manuel complet client"),
        ("Date", date.today().strftime("%d/%m/%Y")),
        ("Depot audite", "github.com/tobiags/Systeme_io_whatsapp_automatisation"),
        ("Revision auditee", "origin/master == local HEAD 8498f5a"),
        ("Mode de lecture", "Suivre les procedures dans l'ordre, puis utiliser les checklists"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.7, 4.8])
    for i, (label, value) in enumerate(meta):
        if i:
            table.add_row()
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], COLORS["light_blue"])
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
            set_cell_margins(cell)

    doc.add_paragraph()
    add_callout(
        doc,
        "Principe du manuel",
        "Chaque fonctionnalite est expliquee avec la meme methode de scaffolding : objectif, prerequis, action, resultat attendu, controle et erreurs a eviter. Le client peut donc utiliser le document comme un mode operatoire, pas seulement comme une description.",
        "EAF7F0",
    )
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.8])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["navy"]
    r.font.size = Pt(10.5)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(9.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.8, 4.7)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, list(widths))
    for i, (label, value) in enumerate(rows):
        if i:
            table.add_row()
            set_table_width(table, list(widths))
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], COLORS["light_blue"])
        for cell in cells:
            set_cell_margins(cell)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        hdr.cells[i].text = header
        set_cell_shading(hdr.cells[i], COLORS["light_blue"])
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
        set_table_width(table, widths)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        add_callout(doc, "Capture manquante", f"La capture attendue n'a pas ete trouvee : {path}", COLORS["soft_yellow"])
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = COLORS["muted"]


def scaffolding_block(
    doc: Document,
    title: str,
    objectif: str,
    prerequis: list[str],
    actions: list[str],
    resultat: str,
    controle: list[str],
    erreurs: list[str] | None = None,
) -> None:
    doc.add_heading(title, level=3)
    add_kv_table(
        doc,
        [
            ("Objectif", objectif),
            ("Prerequis", "\n".join(prerequis) if prerequis else "Aucun prerequis specifique."),
            ("Resultat attendu", resultat),
        ],
        widths=(1.55, 4.95),
    )
    doc.add_paragraph("Methode pas a pas :").runs[0].bold = True
    add_numbers(doc, actions)
    doc.add_paragraph("Controle rapide :").runs[0].bold = True
    add_bullets(doc, controle)
    if erreurs:
        add_callout(doc, "Erreurs a eviter", "\n".join(f"- {e}" for e in erreurs), COLORS["soft_yellow"])


def build_manual() -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_heading("Sommaire", level=1)
    add_numbers(
        doc,
        [
            "Lire le systeme en une page",
            "Acces, liens et roles",
            "Parcours complet d'un lead",
            "Console Admin : chaque element de la page",
            "Portail PILOTAGE LIVE : chaque champ et bouton",
            "Wati, Systeme.io, n8n et Coolify : quoi verifier",
            "Procedures avant, pendant et apres une edition",
            "Broadcasts, templates et incidents connus",
            "Comparaison cahier des charges vs implementation",
            "Annexes : endpoints, templates, scoring, checklists",
        ],
    )

    doc.add_heading("1. Lire le systeme en une page", level=1)
    doc.add_paragraph(
        "La plateforme automatise la relation WhatsApp autour du Challenge Amazon FBA. Elle capte les leads depuis Systeme.io, les range dans une cohorte, les inscrit dans une edition, envoie les messages WhatsApp via Wati, puis adapte les relances selon les informations StreamYard saisies par l'operateur."
    )
    add_callout(
        doc,
        "Regle simple",
        "Systeme.io fait entrer les leads. Wati envoie et recoit les messages. PILOTAGE LIVE alimente la plateforme avec les donnees StreamYard. La Console Admin sert a surveiller.",
        COLORS["soft_green"],
    )
    add_matrix(
        doc,
        ["Outil", "Utilisation client", "Ce qui est automatique", "Ce qui reste manuel"],
        [
            ["Systeme.io", "Capturer les inscriptions", "Webhook vers la plateforme", "Verifier que le funnel et l'automatisation sont actifs"],
            ["Wati", "Lire/repondre aux conversations", "Envoi des templates et reception des webhooks", "Verifier les templates et traiter les reponses humaines"],
            ["Console Admin", "Surveiller les KPIs et la file humaine", "Rafraichissement auto toutes les 60 s", "Analyser et agir si une alerte apparait"],
            ["PILOTAGE LIVE", "Renseigner StreamYard", "Enregistrement en base via API", "Saisir la bonne edition, les liens, inscrits et presents StreamYard"],
            ["n8n", "Orchestration invisible", "Transfert Systeme.io vers API", "Ne pas modifier sans controle technique"],
        ],
        [1.25, 1.75, 1.85, 1.65],
    )

    doc.add_heading("2. Acces, liens et roles", level=1)
    add_kv_table(
        doc,
        [
            ("Console Admin", "http://whatsapp.178.104.229.163.nip.io:3001/"),
            ("Portail PILOTAGE LIVE", "http://whatsapp.178.104.229.163.nip.io:3001/ops/streamyard?token=ops_streamyard_2026_fba_client_x9K4mP2qL7zR"),
            ("API plateforme", "http://whatsapp.178.104.229.163.nip.io"),
            ("Webhook Systeme.io", "POST /webhooks/systemeio"),
            ("Webhook Wati", "POST /webhooks/wati"),
            ("Formulaire closer", "https://www.ecommercecentrale.com/formulaire-challenge"),
            ("Depot GitHub audite", "https://github.com/tobiags/Systeme_io_whatsapp_automatisation"),
        ],
    )
    add_callout(
        doc,
        "Confidentialite",
        "La cle PLATFORM_API_KEY et le token du portail ops donnent acces a des fonctions sensibles. Ils ne doivent pas etre envoyes dans un groupe public ni ajoutes dans une capture partagee.",
        COLORS["soft_yellow"],
    )
    add_matrix(
        doc,
        ["Role", "Responsabilites", "Pages utilisees"],
        [
            ["Client / equipe live", "Creer les lives, saisir les liens, importer inscrits et presents, surveiller Wati.", "PILOTAGE LIVE, Wati"],
            ["Closer", "Reprendre les conversations chaudes et traiter les objections ou intentions d'achat.", "Wati, Console Admin"],
            ["Administrateur", "Verifier les KPIs, les relances humaines, les cohortes, les incidents.", "Console Admin"],
            ["Support technique", "Verifier API, Wati, Celery, base, templates et deploiement.", "API, Coolify, logs, GitHub"],
        ],
        [1.4, 3.4, 1.7],
    )

    doc.add_heading("3. Parcours complet d'un lead", level=1)
    add_numbers(
        doc,
        [
            "Le lead s'inscrit sur un funnel Systeme.io EU ou US-CA.",
            "Systeme.io declenche le webhook, souvent via n8n, vers l'API plateforme.",
            "La plateforme cree ou met a jour le contact, enregistre le consentement et cherche l'edition active de la cohorte.",
            "Le contact est auto-enrole au bon current_step selon le nombre de jours restants avant le challenge.",
            "Le message welcome est envoye via Wati.",
            "Les broadcasts quotidiens envoient l'etape courante et avancent le contact uniquement si l'envoi n'est pas failed.",
            "Autour des lives, l'operateur renseigne les liens, inscrits et presents dans PILOTAGE LIVE.",
            "Les relances J2, J3 et post-challenge changent selon present, inscrit absent ou non inscrit.",
            "Si un achat est detecte avec paid_offer, les messages de campagne s'arretent pour ce contact.",
        ],
    )
    add_matrix(
        doc,
        ["Etat", "Template / action", "Source de decision"],
        [
            ["WELCOME", "welcome", "Inscription Systeme.io"],
            ["COUNTDOWN_J6 a COUNTDOWN_J1", "countdown_j6 a countdown_j1", "Smart skip selon date d'inscription"],
            ["DAY_1", "live_day1", "Edition + lien StreamYard jour 1"],
            ["DAY_2", "live_day2_attended_v2 / registered_absent / not_registered", "Evenements StreamYard jour 1"],
            ["DAY_3", "live_day3_attended_v2 / registered_absent / not_registered", "Evenements StreamYard jour 2"],
            ["AFTER_1", "post_recap_attended / registered_absent / not_registered", "Evenements StreamYard jour 3"],
            ["AFTER_2", "post_testimonials", "Suite post-challenge"],
            ["AFTER_3", "post_inaction_reason", "Suite post-challenge"],
            ["AFTER_4", "post_closer_call", "Lien closer"],
        ],
        [1.5, 3.0, 2.0],
    )

    doc.add_heading("4. Console Admin", level=1)
    add_image(doc, SCREENSHOTS / "admin-dashboard.png", "Console Admin - vue d'ensemble et file de relances humaines")
    scaffolding_block(
        doc,
        "4.1 Page de connexion",
        "Entrer dans la Console Admin de facon securisee.",
        ["Avoir la valeur PLATFORM_API_KEY configuree dans Coolify.", "Ouvrir l'URL de la Console Admin."],
        ["Saisir la cle API dans le champ 'Cle API'.", "Cliquer sur 'Acceder au dashboard'.", "Si la cle est correcte, le dashboard s'ouvre et la cle reste stockee dans le navigateur."],
        "Le tableau de bord s'affiche avec l'indicateur 'En ligne'.",
        ["Si le champ reste vide, le message 'Cle requise' apparait.", "Si l'API refuse la cle, les blocs du dashboard affichent une erreur HTTP."],
        ["Partager la cle API dans une capture.", "Confondre la cle API avec le token du portail StreamYard."],
    )
    scaffolding_block(
        doc,
        "4.2 Bandeau haut",
        "Verifier rapidement que l'on est dans la bonne application et pouvoir se deconnecter.",
        ["Etre connecte a la Console Admin."],
        ["Lire le nom 'Challenge Amazon FBA - Console Admin'.", "Verifier le statut 'En ligne'.", "Utiliser 'Deconnexion' pour retirer la cle du navigateur."],
        "L'utilisateur sait que la console est accessible et peut sortir proprement.",
        ["Le bouton Deconnexion ramene a la page de cle API.", "L'indicateur en ligne signale seulement que l'interface est chargee, pas que Wati accepte les messages."],
    )
    scaffolding_block(
        doc,
        "4.3 Bandeau edition active",
        "Voir l'edition la plus recente connue par la plateforme.",
        ["Avoir au moins une edition enregistree dans la base."],
        ["Lire edition_key, cohorte et date.", "Cliquer sur 'Lien StreamYard' si un streamyard_url global existe."],
        "L'equipe identifie l'edition affichee et peut detecter une edition incoherente.",
        ["L'edition active est l'edition la plus recente en base, pas forcement celle que le client a en tete.", "Si une edition_key a ete mal saisie dans PILOTAGE LIVE, elle peut apparaitre ici."],
        ["Utiliser une edition_key libre comme 'L'OPPORTUNITE AMAZON FBA'. Toujours utiliser un format stable comme 2026-05-21-usca."],
    )
    add_matrix(
        doc,
        ["Bloc", "Ce qu'il mesure", "Comment l'utiliser"],
        [
            ["Contacts", "Nombre total de contacts crees.", "Verifier que les inscriptions Systeme.io alimentent bien la base. Ce nombre n'est pas le nombre eligible broadcast."],
            ["Inscrits campagne", "Nombre de contacts rattaches a une edition/parcours.", "C'est ce compteur qui indique qui peut recevoir les broadcasts automatiques."],
            ["Alerte contacts sans inscription campagne", "Contacts en base non rattaches a une edition.", "Corriger l'edition active ou demander au support de rattacher les contacts opt-in a la bonne edition."],
            ["Messages envoyes", "Nombre de lignes messages, tous statuts confondus.", "Surveiller la progression globale, mais ne pas confondre avec messages delivres."],
            ["Relances humaines", "Messages entrants needs_human.", "Prioriser les conversations a reprendre dans Wati."],
            ["Conversion", "Part de contacts ayant l'evenement paid_offer.", "Suivre la transformation commerciale."],
            ["Presence live par jour", "Contacts avec day1/day2/day3_live_joined.", "Mesurer l'assiduite aux lives."],
            ["Segments", "Froid, tiede, chaud, tres chaud selon score.", "Identifier les leads engageants."],
            ["FAQ detectees", "Intentions FAQ les plus frequentes.", "Comprendre les frictions du public."],
            ["Objections financieres", "Objections detectees dans les messages.", "Aider l'equipe closing a preparer ses reponses."],
            ["Cohortes", "Repartition EU / US-CA.", "Verifier l'equilibre des inscriptions."],
            ["File de relances humaines", "Liste des messages a reprendre.", "Ouvrir Wati et repondre vite aux priorites haute et moyenne."],
        ],
        [1.65, 2.25, 2.6],
    )

    doc.add_heading("5. Portail PILOTAGE LIVE", level=1)
    add_image(doc, SCREENSHOTS / "ops-streamyard-portal.png", "Portail PILOTAGE LIVE - saisie StreamYard sans SSH ni curl")
    scaffolding_block(
        doc,
        "5.1 Acces au portail",
        "Ouvrir la page qui permet de saisir les donnees StreamYard.",
        ["Avoir le lien contenant token=ops_streamyard_2026_fba_client_x9K4mP2qL7zR.", "Utiliser un navigateur web sur ordinateur ou mobile."],
        ["Ouvrir l'URL du portail.", "Verifier que le titre 'PILOTAGE LIVE' apparait.", "Si une page d'erreur indique un token manquant, reprendre le lien complet."],
        "Le portail affiche les sections Contexte du live, Avant le live, Liens de vente et replay, Inscrits, Presents.",
        ["Le token est dans l'URL.", "La page ne demande pas PLATFORM_API_KEY."],
        ["Partager le lien publiquement.", "Ouvrir /ops/streamyard sans le parametre token."],
    )
    scaffolding_block(
        doc,
        "5.2 Contexte du live",
        "Definir a quelle cohorte, edition et jour les prochaines actions seront rattachees.",
        ["Connaitre la cohorte : EU ou US-CA.", "Connaitre l'edition_key officielle.", "Savoir si l'action concerne le jour 1, 2 ou 3."],
        ["Choisir Cohorte : US/CA ou EU. Si le challenge n'a qu'une cohorte USA, choisir toujours US/CA.", "Saisir edition_key au format obligatoire : AAAA-MM-JJ-usca ou AAAA-MM-JJ-eu. Exemple reel : 2026-05-28-usca.", "Choisir Jour 1, Jour 2 ou Jour 3 selon le lien ou la liste que l'on va enregistrer."],
        "Les actions suivantes enregistrent les donnees sur la bonne edition.",
        ["Le meme edition_key doit etre utilise pour tous les liens, inscrits et presents de la meme edition.", "Changer de jour modifie uniquement l'action de live/inscrits/presents, pas les liens commerciaux."],
        ["Saisir un titre commercial a la place de l'edition_key, par exemple 'L'OPPORTUNITE AMAZON FBA'.", "Melanger EU et US-CA avec la meme edition.", "Changer l'edition_key entre jour 1, jour 2 et jour 3."],
    )
    add_callout(
        doc,
        "Exemple concret pour une cohorte USA",
        "Si le challenge commence le 28/05/2026 et qu'il n'y a qu'une cohorte USA, utiliser Cohorte = US/CA et Edition key = 2026-05-28-usca partout : live jour 1, live jour 2, live jour 3, replays, lien paiement, inscrits StreamYard et presents StreamYard.",
        COLORS["soft_yellow"],
    )
    scaffolding_block(
        doc,
        "5.3 Avant le live - Enregistrer le live",
        "Associer le lien StreamYard du jour a l'edition et a la cohorte.",
        ["Avoir cree le live StreamYard du jour.", "Avoir copie le lien StreamYard exact."],
        ["Renseigner le contexte du live.", "Choisir Jour 1, coller le lien jour 1, puis cliquer sur 'Enregistrer le live'.", "Choisir Jour 2, remplacer par le lien jour 2, puis cliquer sur 'Enregistrer le live'.", "Choisir Jour 3, remplacer par le lien jour 3, puis cliquer sur 'Enregistrer le live'.", "Lire le message de succes apres chaque jour."],
        "Les rappels live du jour utiliseront ce lien dans les templates live_day*.",
        ["Le message de succes mentionne la cohorte et le jour.", "En base, le lien alimente day1_url, day2_url ou day3_url."],
        ["Enregistrer uniquement le jour 1 et oublier les jours 2 et 3.", "Coller une URL de page admin StreamYard au lieu du lien participant."],
    )
    scaffolding_block(
        doc,
        "5.4 Liens de vente et replay",
        "Enregistrer les URLs variables utilisees dans l'offre et les messages post-challenge.",
        ["Connaitre le lien paiement.", "Connaitre le lien closer / reservation.", "Avoir les trois replays ou la page replay."],
        ["Renseigner le contexte avec la bonne edition_key.", "Remplir Lien paiement.", "Remplir Lien closer / reservation.", "Remplir Replay jour 1, 2 et 3.", "Cliquer sur 'Enregistrer les liens'."],
        "Les messages d'offre, closer et replay utiliseront les URLs rattachees a cette edition.",
        ["Le message de succes confirme l'enregistrement.", "Revenir a la meme edition_key pour corriger un lien si besoin."],
        ["Creer une nouvelle edition en saisissant une mauvaise edition_key.", "Laisser les replay vides alors que le template Wati attend des URLs."],
    )
    scaffolding_block(
        doc,
        "5.5 Juste avant / debut du live - Envoyer les inscrits",
        "Dire a la plateforme qui s'est inscrit sur la page StreamYard.",
        ["Exporter ou copier les inscrits StreamYard.", "S'assurer que les numeros existent deja comme contacts."],
        ["Choisir 'Coller les numeros' ou 'Importer un CSV'.", "Coller un numero par ligne ou selectionner le CSV.", "Verifier le nombre de numeros detectes.", "Cliquer sur 'Envoyer les inscrits'."],
        "La plateforme cree des evenements dayN_streamyard_registered.",
        ["Le message indique enregistres, deja connus et non trouves.", "Les inscrits alimentent la branche 'registered_absent' si la personne ne vient pas au live."],
        ["Envoyer les presents dans cette section.", "Importer un CSV qui ne contient pas les numeros de telephone.", "Utiliser cette section pour rattacher les contacts Systeme.io au parcours. Les contacts Systeme.io sont geres par l'inscription campagne, pas par cette zone."],
    )
    add_callout(
        doc,
        "Attention : Systeme.io et StreamYard ne sont pas la meme liste",
        "Les blocs 'Envoyer les inscrits' et 'Envoyer les presents' concernent uniquement les listes exportees depuis StreamYard. Ils ne servent pas a ajouter les contacts Systeme.io dans la sequence broadcast. Si la zone affiche 'Numeros detectes : 0', cela signifie qu'aucun vrai numero n'a ete colle ou importe; le texte visible dans la zone est seulement un exemple.",
        COLORS["soft_yellow"],
    )
    scaffolding_block(
        doc,
        "5.6 Apres le live - Envoyer les presents",
        "Dire a la plateforme qui a reellement assiste au live.",
        ["Avoir la liste des participants presents au live.", "Avoir selectionne la bonne cohorte, edition et jour."],
        ["Choisir 'Coller les numeros' ou 'Importer un CSV'.", "Coller ou importer les numeros.", "Verifier le nombre detecte.", "Cliquer sur 'Envoyer les presents'."],
        "La plateforme cree des evenements dayN_live_joined et les prochaines relances utilisent la branche present.",
        ["Le message indique presents enregistres, deja connus et non trouves.", "Les presents jour 1 influencent DAY_2; les presents jour 2 influencent DAY_3; les presents jour 3 influencent AFTER_1."],
        ["Envoyer les inscrits a la place des presents.", "Choisir le mauvais jour avant d'envoyer."],
    )

    doc.add_heading("6. Wati, Systeme.io, n8n et Coolify", level=1)
    doc.add_heading("6.1 Wati", level=2)
    scaffolding_block(
        doc,
        "Utiliser Wati au quotidien",
        "Lire et reprendre les conversations WhatsApp des prospects.",
        ["Avoir acces au workspace Wati.", "Verifier que le numero WhatsApp Business est connecte."],
        ["Ouvrir l'inbox Wati.", "Surveiller les nouvelles reponses.", "Traiter en priorite les demandes d'appel, objections fortes et problemes de paiement.", "Verifier les templates avant chaque edition."],
        "Les prospects recoivent les messages et l'equipe peut repondre manuellement.",
        ["Template approuve.", "Categorie compatible avec le pays cible.", "Variables du template alignees avec le code."],
        ["Modifier un template Wati sans adapter le code.", "Ignorer les restrictions Meta sur les templates marketing vers les destinataires US."],
    )
    add_callout(
        doc,
        "Point d'attention broadcast",
        "Un message peut etre cree dans la table messages avec le statut failed. La Console Admin compte les messages tous statuts confondus. Pour diagnostiquer une panne de reception, il faut verifier Wati et le statut provider, pas seulement le nombre total.",
        COLORS["soft_yellow"],
    )
    add_callout(
        doc,
        "Restriction connue : numeros +1 et templates MARKETING",
        "Depuis la restriction Meta sur les templates marketing vers les numeros US (+1), un template Wati approuve en categorie MARKETING peut etre refuse pour certains destinataires +1. Dans ce cas, le systeme conserve le contact au meme current_step afin d'eviter une progression artificielle. La correction durable consiste a creer des rappels live en categorie UTILITY lorsque le contenu le permet.",
        COLORS["soft_red"],
    )
    doc.add_heading("6.2 Systeme.io et n8n", level=2)
    add_image(doc, SCREENSHOTS / "n8n-workflows.png", "n8n - workflows d'orchestration entre Systeme.io et la plateforme")
    scaffolding_block(
        doc,
        "Verifier le funnel d'inscription",
        "S'assurer que chaque inscription arrive dans la plateforme.",
        ["Avoir acces au funnel Systeme.io.", "Connaitre la cohorte du funnel : EU ou US-CA."],
        ["Ouvrir l'automatisation du funnel.", "Verifier que le webhook vers la plateforme ou n8n est actif.", "Faire une inscription test.", "Verifier dans Wati que le welcome part, puis dans la Console Admin que le contact augmente."],
        "Le lead est cree, consentement enregistre, enrollment cree, welcome envoye.",
        ["Le payload doit permettre d'identifier la cohorte.", "Une edition future active doit exister pour que l'auto-enrollment fonctionne."],
        ["Tester sans edition future active.", "Changer la structure du webhook Systeme.io sans verifier la normalisation."],
    )
    doc.add_heading("6.3 Coolify et configuration", level=2)
    add_image(doc, SCREENSHOTS / "coolify-secrets.png", "Coolify - configuration des secrets de production")
    add_matrix(
        doc,
        ["Variable", "Utilisation", "Qui doit la modifier"],
        [
            ["POSTGRES_DSN", "Connexion base PostgreSQL.", "Support technique uniquement"],
            ["REDIS_URL", "Celery worker et beat.", "Support technique uniquement"],
            ["WATI_API_URL", "Endpoint tenant Wati.", "Support technique"],
            ["WATI_API_TOKEN", "Authentification Wati.", "Support technique"],
            ["PLATFORM_API_KEY", "Acces Console Admin/API protegee.", "Administrateur"],
            ["OPS_PORTAL_TOKEN", "Acces PILOTAGE LIVE.", "Administrateur"],
            ["PROGRAM_PAYMENT_URL", "Fallback lien paiement si edition vide.", "Support technique / client"],
            ["ONCEHUB_FORM_URL", "Fallback lien formulaire closer.", "Support technique / client"],
            ["SMTP_*", "Notifications email closers.", "Support technique"],
        ],
        [1.6, 3.0, 1.9],
    )

    doc.add_heading("7. Procedures d'exploitation", level=1)
    doc.add_heading("7.1 Avant chaque edition", level=2)
    add_numbers(
        doc,
        [
            "Verifier que Wati est connecte et que les templates necessaires sont approuves.",
            "Verifier la categorie Wati des templates live. Pour les destinataires +1, privilegier des templates UTILITY strictement informatifs si Meta/Wati les approuve.",
            "Creer ou verifier les lives StreamYard jour 1, jour 2 et jour 3 pour chaque cohorte active.",
            "Definir l'edition_key officielle et l'utiliser partout sans variation.",
            "Ouvrir PILOTAGE LIVE et enregistrer les liens StreamYard par jour.",
            "Enregistrer le lien paiement, le lien closer et les liens replay si deja disponibles.",
            "Faire une inscription test Systeme.io pour verifier contact, consentement, enrollment et welcome.",
            "Surveiller la Console Admin pour confirmer que les compteurs evoluent normalement.",
        ],
    )
    doc.add_heading("7.2 Pendant chaque live", level=2)
    add_numbers(
        doc,
        [
            "Avant le live ou au debut, envoyer les inscrits StreamYard via la section 'Envoyer les inscrits'.",
            "Pendant le jour 3, declencher le message d'offre H+2 si l'equipe l'utilise. Cette fonction existe via API; si le client veut un bouton, il faut l'ajouter au portail.",
            "Lire les reponses importantes dans Wati.",
            "Apres le live, envoyer les presents via la section 'Envoyer les presents'.",
            "Verifier les messages de succes et noter les numeros non trouves.",
        ],
    )
    doc.add_heading("7.3 Apres le challenge", level=2)
    add_numbers(
        doc,
        [
            "Verifier que les replays et le lien closer sont bien renseignes sur la bonne edition_key.",
            "Surveiller les relances post-challenge et les reponses entrantes.",
            "Traiter la file humaine dans Wati et la Console Admin.",
            "Verifier les objections financieres et les intentions d'achat.",
            "Exporter ou noter les apprentissages pour ameliorer la prochaine edition.",
        ],
    )

    doc.add_heading("8. Broadcasts, templates et incidents", level=1)
    add_callout(
        doc,
        "Ce que signifie broadcast dans ce projet",
        "Le broadcast n'est pas un message envoye dans un groupe WhatsApp. L'API WhatsApp Business envoie des messages 1-to-1 a chaque contact inscrit, via des templates Wati.",
        COLORS["soft_green"],
    )
    add_matrix(
        doc,
        ["Situation", "Cause probable", "Action minimale"],
        [
            ["Personne ne recoit plus les messages", "Provider Wati retourne failed, template non aligne, restriction Meta ou edition mal renseignee.", "Tester un seul envoi Wati, verifier le statut messages, ne pas avancer les contacts manuellement."],
            ["Mauvais lien live", "day1/day2/day3_url mal saisi ou mauvaise edition_key.", "Re-saisir le lien sur la bonne edition et le bon jour dans PILOTAGE LIVE."],
            ["Mauvaise branche present/absent", "Inscrits ou presents non envoyes, mauvais jour selectionne.", "Envoyer les bonnes listes StreamYard avec le bon jour."],
            ["Welcome ne part pas", "Template Wati, restriction Meta, contact invalide, token Wati.", "Tester le template welcome sur un numero valide et verifier Wati."],
            ["US/CA bloque", "Meta peut restreindre les templates marketing vers des destinataires US.", "Verifier categorie du template et envisager un template utility si admissible."],
            ["Contacts avancent mal", "Ancienne version de prod ou relance manuelle non ciblee.", "Verifier que le code ne progresse plus sur status failed et relancer seulement l'edition concernee."],
        ],
        [1.7, 2.45, 2.35],
    )
    doc.add_heading("8.1 Methode Utility pour les rappels live US", level=2)
    add_callout(
        doc,
        "Objectif",
        "Permettre aux numeros US/CA (+1) de recevoir les informations pratiques du live lorsque les templates marketing sont refuses par Meta/Wati. La methode ne consiste pas a contourner Meta : elle consiste a soumettre des messages reellement transactionnels/informatifs, sans promesse commerciale.",
        COLORS["soft_green"],
    )
    add_matrix(
        doc,
        ["Regle", "A faire", "A eviter"],
        [
            ["Intention", "Confirmer l'acces a un live auquel la personne s'est inscrite.", "Vendre, convaincre, relancer commercialement."],
            ["Texte", "Formulation neutre : lien, heure, rappel pratique.", "Mots comme opportunite, offre, dernier moment, gagnez, transformez votre business."],
            ["Variables", "{{1}} prenom, {{2}} lien du live, {{3}} heure.", "Ajouter un lien de paiement ou une promesse de resultat dans le rappel live."],
            ["Ciblage", "Utiliser Utility en priorite pour les numeros +1 si Wati l'approuve.", "Envoyer en masse un template MARKETING aux +1 en esperant que Meta accepte."],
            ["Suivi", "Relancer uniquement les contacts restes au step bloque apres failed.", "Rejouer un broadcast global aux contacts deja passes au step suivant."],
        ],
        [1.35, 2.75, 2.4],
    )
    add_callout(
        doc,
        "Exemple de template Utility propose",
        "Bonjour {{1}}, voici votre acces au live du Challenge Amazon FBA auquel vous etes inscrit. Lien : {{2}}. Heure : {{3}}. Ce message est un rappel pratique lie a votre inscription.",
        COLORS["light_gray"],
    )
    add_callout(
        doc,
        "Regle de securite operationnelle",
        "Ne jamais relancer un broadcast global si on ne sait pas si l'echec est total ou partiel. Une relance globale peut envoyer l'etape suivante a ceux qui avaient deja recu l'etape precedente.",
        COLORS["soft_red"],
    )

    doc.add_heading("9. Comparaison cahier des charges vs implementation", level=1)
    doc.add_paragraph(
        "Le cahier des charges V2 a ete compare au code du depot GitHub tobiags/Systeme_io_whatsapp_automatisation. Le depot distant origin/master correspond au HEAD local audite : 6239425."
    )
    add_matrix(
        doc,
        ["Exigence du cahier", "Etat implementation", "Commentaire client"],
        [
            ["Per-day URLs day1/day2/day3", "Implemente", "Saisie via PILOTAGE LIVE; attention a l'edition_key."],
            ["Auto-enrollment Systeme.io", "Implemente", "Necessite une edition future active pour la cohorte."],
            ["Smart skip J-6 a J0", "Implemente", "Le contact commence a la bonne etape selon la date d'inscription."],
            ["Branching 3 voies", "Implemente", "Depend des listes inscrits et presents envoyees via StreamYard."],
            ["Message H+2 Jour 3", "Implemente cote API/Celery", "Pas de bouton visible dans le portail; declenchement technique possible."],
            ["OnceHub post-challenge", "Implemente", "Lien par edition ou fallback settings.oncehub_form_url."],
            ["Notifications email closers", "Implemente avec degradation", "Fonctionne seulement si SMTP et emails sont configures."],
            ["Dashboard admin", "Implemente", "KPIs, segments, FAQ, objections, file humaine."],
            ["Multi-utilisateurs dashboard", "Non implemente", "Un acces admin simple par cle API."],
            ["Broadcast groupe WhatsApp", "Impossible techniquement", "WhatsApp Business API fonctionne en 1-to-1."],
            ["Automatisation StreamYard complete", "Partiellement manuel", "Le portail remplace Termius/curl, mais StreamYard n'a pas d'API publique exploitee ici."],
            ["Alignement templates Wati", "A surveiller", "Les noms et variables Wati doivent rester strictement alignes avec le code."],
        ],
        [2.0, 1.7, 2.8],
    )

    doc.add_heading("10. Annexes", level=1)
    doc.add_heading("10.1 Endpoints utiles", level=2)
    add_matrix(
        doc,
        ["Endpoint", "Role", "Usage client"],
        [
            ["GET /health", "Verifier que l'API et la DB repondent.", "Support / controle simple"],
            ["GET /dashboard/summary", "Donnees de la Console Admin.", "Automatique via front"],
            ["GET /webhooks/wati/queue", "File humaine.", "Automatique via front"],
            ["POST /webhooks/systemeio", "Entree inscriptions.", "Configure dans Systeme.io/n8n"],
            ["POST /webhooks/wati", "Messages entrants Wati.", "Configure dans Wati"],
            ["POST /ops/streamyard/session", "Lien live du jour.", "Bouton Enregistrer le live"],
            ["POST /ops/streamyard/resources", "Liens paiement/closer/replays.", "Bouton Enregistrer les liens"],
            ["POST /ops/streamyard/registrants", "Inscrits StreamYard.", "Bouton Envoyer les inscrits"],
            ["POST /ops/streamyard/attendance", "Presents live.", "Bouton Envoyer les presents"],
            ["POST /campaigns/broadcast", "Relance manuelle ciblee.", "Support technique uniquement"],
            ["POST /campaigns/trigger/day3-offer", "Offre H+2 Jour 3.", "Support technique ou futur bouton"],
        ],
        [2.0, 2.65, 1.85],
    )
    doc.add_heading("10.2 Scoring et segments", level=2)
    add_matrix(
        doc,
        ["Signal", "Points", "Interpretation"],
        [
            ["registered", "10", "Inscription initiale"],
            ["group_whatsapp_joined", "15", "A rejoint le groupe"],
            ["opened_message", "5", "A ouvert un message"],
            ["clicked_link / streamyard_link_clicked", "10", "Clique un lien"],
            ["replied_message / poll_answered", "10", "Interagit par reponse"],
            ["dayN_streamyard_registered", "5", "Inscrit sur StreamYard"],
            ["day1_live_joined", "30", "Present live jour 1"],
            ["day2/day3_live_joined", "25", "Present live jour 2 ou 3"],
            ["asked_question", "20", "Question explicite"],
            ["conversion_intent_detected", "35", "Signal fort d'achat"],
            ["paid_offer", "50", "Achat detecte, stop broadcasts"],
        ],
        [2.55, 0.8, 3.15],
    )
    add_bullets(doc, ["Score <= 15 : froid", "Score <= 40 : tiede", "Score <= 75 : chaud", "Score > 75 : tres chaud"])

    doc.add_heading("10.3 Checklist finale client", level=2)
    add_matrix(
        doc,
        ["Moment", "Checklist"],
        [
            ["Avant edition", "Wati connecte; templates approuves; edition_key definie; lives StreamYard crees; liens saisis; inscription test faite."],
            ["Avant chaque live", "Bon jour choisi; bon lien StreamYard enregistre; inscrits envoyes."],
            ["Apres chaque live", "Presents envoyes; Wati surveille; file humaine traitee."],
            ["Apres challenge", "Replays/closer verifies; objections traitees; incidents notes pour prochaine edition."],
            ["En cas d'incident", "Ne pas relancer globalement; tester un contact; verifier Wati/template/status; relancer cible uniquement."],
        ],
        [1.35, 5.15],
    )

    doc.add_paragraph()
    add_callout(
        doc,
        "Conclusion",
        "Le systeme est exploitable sans intervention technique quotidienne si trois disciplines sont respectees : utiliser toujours la bonne edition_key, garder les templates Wati alignes avec le code, et renseigner StreamYard dans PILOTAGE LIVE avant/apres chaque live.",
        COLORS["soft_green"],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_manual()
