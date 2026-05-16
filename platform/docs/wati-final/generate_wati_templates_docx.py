from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "templates_wati_final.json"
OUTPUT = ROOT / "TEMPLATES-WATI-FINAL.docx"

ACCENT = RGBColor(0x1D, 0x4E, 0x89)
ACCENT_LIGHT = "D9E8FB"
BORDER = "D1D5DB"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    for level, size in [(1, 20), (2, 15), (3, 12)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)


def add_cover(doc: Document, meta: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Templates Wati Finalises")
    run.font.name = "Aptos Display"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Challenge Amazon FBA")
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Voix retenue : le formateur ecrit directement")
    run.font.size = Pt(11)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Projet", meta["project"]),
        ("Strategie", meta["strategy"]),
        ("Categorie", f'{meta["category"]} / {meta["subCategory"]} / {meta["language"]}'),
        ("Genere le", datetime.now().strftime("%d/%m/%Y a %H:%M")),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], ACCENT_LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(table)
    doc.add_page_break()


def add_meta(doc: Document, meta: dict) -> None:
    doc.add_paragraph("Positionnement", style="Heading 1")
    for note in meta["notes"]:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.add_run("• " + note)


def add_template(doc: Document, template: dict) -> None:
    doc.add_paragraph(template["key"], style="Heading 2")

    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Nom Wati", template["name"]),
        ("Phase", template["phase"]),
        ("Variables", ", ".join(f'{{{{{v["name"]}}}}}' for v in template["variables"])),
    ]
    for row, (label, value) in zip(info.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], ACCENT_LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(info)

    doc.add_paragraph("Exemples de variables", style="Heading 3")
    var_table = doc.add_table(rows=1, cols=3)
    var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Variable", "Usage", "Exemple"]
    for i, text in enumerate(headers):
        var_table.rows[0].cells[i].text = text
        shade_cell(var_table.rows[0].cells[i], ACCENT_LIGHT)
        var_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for var in template["variables"]:
        row = var_table.add_row().cells
        row[0].text = f'{{{{{var["name"]}}}}}'
        row[1].text = var["description"]
        row[2].text = var["sample"]
    set_table_borders(var_table)

    doc.add_paragraph("Texte final", style="Heading 3")
    for line in template["body"].split("\n"):
        doc.add_paragraph(line)


def main() -> None:
    payload = json.loads(SOURCE.read_text(encoding="utf-8"))
    doc = Document()
    configure(doc)
    add_cover(doc, payload["meta"])
    add_meta(doc, payload["meta"])
    for template in payload["templates"]:
        add_template(doc, template)
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    main()
