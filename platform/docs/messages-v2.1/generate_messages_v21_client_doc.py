from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE_DOC = Path(r"C:\Users\tobid\Downloads\messages-challenge-amazon-fba-v2.docx")
ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "messages-challenge-amazon-fba-v2.1-validation-client.docx"
OUTPUT_MD = ROOT / "messages-challenge-amazon-fba-v2.1-validation-client.md"

TABLE_HEAD = "D9E8FB"
SUCCESS = "EAF6EE"
INFO = "EAF1FB"
BORDER = "CBD5E1"


@dataclass
class TemplateDelta:
    key: str
    action: str
    timing: str
    audience: str
    variables: list[tuple[str, str]]
    text: list[str]
    note: str | None = None


DELTAS: list[TemplateDelta] = [
    TemplateDelta(
        key="live_day2_attended",
        action="Remplacement cible",
        timing="Jour 2, 2h avant le live",
        audience="Presents au Jour 1",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 2"),
            ("{{3}}", "Heure Session 2"),
        ],
        text=[
            "{{1}}, ta presence d'hier soir montre deja que tu es dans une vraie demarche.",
            "",
            "Ce soir, on va encore plus loin.",
            "",
            "Session 2 a {{3}} - le sujet : Construire Ton Business Amazon Pas a Pas",
            "",
            "📹 {{2}}",
            "",
            "Avant ce soir, dis-moi : qu'est-ce qui t'a le plus marque hier ?",
        ],
        note="On conserve l'esprit du message v2 et on ajoute seulement la question demandee par le client.",
    ),
    TemplateDelta(
        key="live_day3_attended",
        action="Remplacement cible",
        timing="Jour 3, 2h avant le live",
        audience="Presents au Jour 2",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 3"),
            ("{{3}}", "Heure Session 3"),
        ],
        text=[
            "{{1}}, ce soir, on arrive a la derniere session.",
            "",
            "On se retrouve a {{3}} pour : Le Secret des Vendeurs a Succes sur Amazon",
            "",
            "📹 {{2}}",
            "",
            "Et avant ce soir, dis-moi : qu'est-ce qui t'a le plus marque dans la session d'hier ?",
        ],
        note="On garde la base redactionnelle et on ajoute seulement la question contextuelle demandee.",
    ),
    TemplateDelta(
        key="post_recap_registered_absent",
        action="Remplacement cible",
        timing="J+1",
        audience="Inscrits ayant suivi une partie du challenge",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Replay Jour 1"),
            ("{{3}}", "Replay Jour 2"),
            ("{{4}}", "Replay Jour 3"),
        ],
        text=[
            "Bonjour {{1}} 👋",
            "",
            "Comme promis, voici les replays des 3 jours du Challenge Amazon FBA :",
            "",
            "🎥 Jour 1 : {{2}}",
            "🎥 Jour 2 : {{3}}",
            "🎥 Jour 3 : {{4}}",
            "",
            "Prends le temps de regarder ce que tu as manque, puis dis-moi ce qui t'a le plus parle.",
        ],
        note="Ce message remplace la logique 'reponds INFO' par un partage direct des replays.",
    ),
    TemplateDelta(
        key="post_recap_not_registered",
        action="Remplacement cible",
        timing="J+1",
        audience="Aucun live suivi",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Replay Jour 1"),
            ("{{3}}", "Replay Jour 2"),
            ("{{4}}", "Replay Jour 3"),
        ],
        text=[
            "Bonjour {{1}} 👋",
            "",
            "Tu n'as peut-etre pas pu suivre les lives, donc je te laisse ici les replays des 3 jours :",
            "",
            "🎥 Jour 1 : {{2}}",
            "🎥 Jour 2 : {{3}}",
            "🎥 Jour 3 : {{4}}",
            "",
            "Si tu les regardes, reponds-moi simplement ici. Je veux savoir ce qui t'a le plus parle.",
        ],
        note="On garde l'intention de rattrapage, avec un lien concret plutot qu'un CTA flou.",
    ),
    TemplateDelta(
        key="live_day1_h10",
        action="Nouveau message",
        timing="Jour 1, 10 minutes avant",
        audience="Tous les inscrits actifs",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 1"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "On demarre dans 10 minutes.",
            "",
            "Voici ton lien d'acces :",
            "{{2}}",
            "",
            "A tout de suite.",
        ],
    ),
    TemplateDelta(
        key="live_day1_hplus5",
        action="Nouveau message",
        timing="Jour 1, 5 minutes apres le debut",
        audience="Ceux qui n'ont pas clique ou ouvert le message precedent",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 1"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "Le live a commence il y a 5 minutes.",
            "",
            "Si tu comptais nous rejoindre, tu peux encore entrer ici :",
            "{{2}}",
        ],
    ),
    TemplateDelta(
        key="live_day2_h10",
        action="Nouveau message",
        timing="Jour 2, 10 minutes avant",
        audience="Tous les contacts cibles pour le Jour 2",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 2"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "On demarre dans 10 minutes.",
            "",
            "Voici ton lien pour la Session 2 :",
            "{{2}}",
            "",
            "A tout de suite.",
        ],
    ),
    TemplateDelta(
        key="live_day2_hplus5",
        action="Nouveau message",
        timing="Jour 2, 5 minutes apres le debut",
        audience="Ceux qui n'ont pas clique ou ouvert le message precedent",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 2"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "La Session 2 a commence il y a 5 minutes.",
            "",
            "Tu peux encore nous rejoindre ici :",
            "{{2}}",
        ],
    ),
    TemplateDelta(
        key="live_day3_h10",
        action="Nouveau message",
        timing="Jour 3, 10 minutes avant",
        audience="Tous les contacts cibles pour le Jour 3",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 3"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "Plus que 10 minutes avant la derniere session.",
            "",
            "Voici ton lien :",
            "{{2}}",
            "",
            "A tout de suite.",
        ],
    ),
    TemplateDelta(
        key="live_day3_hplus5",
        action="Nouveau message",
        timing="Jour 3, 5 minutes apres le debut",
        audience="Ceux qui n'ont pas clique ou ouvert le message precedent",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien StreamYard Session 3"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "La derniere session a commence il y a 5 minutes.",
            "",
            "Si tu veux encore nous rejoindre, c'est ici :",
            "{{2}}",
        ],
    ),
    TemplateDelta(
        key="live_day3_offer_hplus2",
        action="Nouveau message",
        timing="Jour 3, 2 heures apres le debut du live",
        audience="Contacts n'ayant pas encore paye",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien de paiement"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "Comme promis pendant le direct, voici le lien pour rejoindre l'accompagnement :",
            "{{2}}",
            "",
            "Prends le temps de le lire tranquillement.",
            "",
            "Et si tu veux qu'on echange avant de decider, reponds-moi ici.",
        ],
        note="A arreter automatiquement pour les personnes qui ont deja pris la formation le jour meme.",
    ),
    TemplateDelta(
        key="post_testimonials",
        action="Nouveau message",
        timing="Apres J+1",
        audience="Contacts encore actifs",
        variables=[("{{1}}", "Prenom")],
        text=[
            "Bonjour {{1}},",
            "",
            "Je te partage deux retours de personnes qui sont passees a l'action apres le challenge :",
            "",
            "[TEMOIGNAGE 1]",
            "",
            "[TEMOIGNAGE 2]",
            "",
            "Si tu veux qu'on voie si c'est possible pour toi aussi, reponds-moi ici.",
        ],
        note="A completer avec les vrais temoignages transmis par le client.",
    ),
    TemplateDelta(
        key="post_inaction_reason",
        action="Nouveau message",
        timing="Apres les temoignages",
        audience="Contacts qui n'ont toujours pas avance",
        variables=[("{{1}}", "Prenom")],
        text=[
            "Bonjour {{1}},",
            "",
            "J'aimerais comprendre quelque chose.",
            "",
            "Qu'est-ce qui t'empeche aujourd'hui de passer a l'action : le budget, le temps, la peur de te tromper, ou autre chose ?",
            "",
            "Reponds-moi franchement.",
        ],
    ),
    TemplateDelta(
        key="post_closer_call",
        action="Nouveau message",
        timing="Derniere relance",
        audience="Contacts encore chauds",
        variables=[
            ("{{1}}", "Prenom"),
            ("{{2}}", "Lien de reservation closer"),
        ],
        text=[
            "Bonjour {{1}},",
            "",
            "Si tu veux qu'on regarde ta situation plus serieusement, tu peux reserver un echange ici :",
            "{{2}}",
            "",
            "Tu pourras poser tes questions et voir si l'accompagnement est vraiment fait pour toi.",
        ],
    ),
]


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def upgrade_cover(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if "Version 2.0" in paragraph.text:
            paragraph.text = paragraph.text.replace("Version 2.0", "Version 2.1")
        if "validation v2" in paragraph.text.lower():
            paragraph.text = paragraph.text.replace("validation v2", "validation v2.1")
            paragraph.text = paragraph.text.replace("Validation v2", "Validation v2.1")
        if "18 templates" in paragraph.text:
            paragraph.text = paragraph.text.replace("18 templates", "18 templates de base + ajouts v2.1")


def add_intro_delta(doc: Document) -> None:
    doc.add_page_break()
    doc.add_paragraph("7. Ajouts et corrections demandes - Version 2.1", style="Heading 1")
    bullets = [
        "Base editoriale conservee : les messages existants du document v2 restent la reference.",
        "Objectif : ajouter uniquement les nouveaux rappels, les messages manquants et les corrections explicitement demandees par le client.",
        "Les remplacements ci-dessous concernent seulement les messages pour lesquels le client a demande un ajustement precis.",
    ]
    for bullet in bullets:
        p = doc.add_paragraph()
        p.add_run("• " + bullet)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, head in enumerate(["Type", "Cle Wati", "Action retenue"]):
        table.rows[0].cells[idx].text = head
        shade_cell(table.rows[0].cells[idx], TABLE_HEAD)
        table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

    rows = [
        ("Messages de base", "Messages 1 a 15 + post_followup", "Conserves sur la base du document v2"),
        ("Remplacements cibles", "live_day2_attended / live_day3_attended / post_recap_registered_absent / post_recap_not_registered", "Ajustements demandes par le client"),
        ("Nouveaux messages", "rappels H-10 / H+5, paiement H+2, temoignages, frein, closer", "Ajouts pour completer la sequence"),
    ]
    for left, middle, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = middle
        row[2].text = right
    set_table_borders(table)


def add_template_delta(doc: Document, delta: TemplateDelta) -> None:
    doc.add_paragraph(delta.key, style="Heading 2")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Type", delta.action),
        ("Timing", delta.timing),
        ("Audience", delta.audience),
        ("Variables", ", ".join(key for key, _ in delta.variables)),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], TABLE_HEAD)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(meta)

    doc.add_paragraph("Variables du message", style="Heading 3")
    vars_table = doc.add_table(rows=1, cols=2)
    vars_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    vars_table.rows[0].cells[0].text = "Variable"
    vars_table.rows[0].cells[1].text = "Valeur injectee"
    shade_cell(vars_table.rows[0].cells[0], TABLE_HEAD)
    shade_cell(vars_table.rows[0].cells[1], TABLE_HEAD)
    vars_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    vars_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    for key, value in delta.variables:
        row = vars_table.add_row().cells
        row[0].text = key
        row[1].text = value
    set_table_borders(vars_table)

    doc.add_paragraph("Texte du message WhatsApp", style="Heading 3")
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade_cell(box.rows[0].cells[0], INFO)
    set_table_borders(box)
    cell = box.rows[0].cells[0]
    for line in delta.text:
        cell.add_paragraph(line)
    if cell.paragraphs and not cell.paragraphs[0].text.strip():
        cell._tc.remove(cell.paragraphs[0]._p)

    if delta.note:
        note = doc.add_table(rows=1, cols=1)
        note.alignment = WD_TABLE_ALIGNMENT.CENTER
        note.rows[0].cells[0].text = delta.note
        shade_cell(note.rows[0].cells[0], SUCCESS)
        set_table_borders(note)


def render_docx() -> None:
    doc = Document(SOURCE_DOC)
    upgrade_cover(doc)
    add_intro_delta(doc)
    for delta in DELTAS:
        add_template_delta(doc, delta)
    doc.save(OUTPUT_DOCX)


def render_markdown() -> None:
    lines = [
        "# Challenge Amazon FBA",
        "",
        "## Sequence de messages WhatsApp - Version 2.1",
        "",
        "Base editoriale conservee depuis le document v2.",
        "",
        "Objectif de cette version : ajouter seulement les nouveaux rappels, les remplacements explicitement demandes et les messages manquants.",
        "",
        "## Matrice de travail",
        "",
        "- Messages de base : conserves sur la base du document v2",
        "- Remplacements cibles : `live_day2_attended`, `live_day3_attended`, `post_recap_registered_absent`, `post_recap_not_registered`",
        "- Nouveaux messages : rappels H-10 / H+5, paiement H+2, temoignages, frein, closer",
        "",
    ]
    for delta in DELTAS:
        lines.append(f"## {delta.key}")
        lines.append("")
        lines.append(f"- Type : {delta.action}")
        lines.append(f"- Timing : {delta.timing}")
        lines.append(f"- Audience : {delta.audience}")
        lines.append("")
        lines.append("Variables :")
        for key, desc in delta.variables:
            lines.append(f"- {key} : {desc}")
        lines.append("")
        lines.append("```text")
        lines.extend(delta.text)
        lines.append("```")
        if delta.note:
            lines.append("")
            lines.append(f"Note : {delta.note}")
        lines.append("")
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    render_docx()
    render_markdown()
    print(f"Generated {OUTPUT_DOCX}")
    print(f"Generated {OUTPUT_MD}")


if __name__ == "__main__":
    main()
