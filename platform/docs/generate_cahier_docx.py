from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "CAHIER-DES-CHARGES-V2.md"
OUTPUT = ROOT / "platform" / "docs" / "CAHIER-DES-CHARGES-V2.docx"


ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_LIGHT = "D9EAF7"
GRAY_LIGHT = "F3F5F7"
BORDER = "C9D3DD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clean_inline(text: str) -> str:
    text = text.replace("â€”", "—").replace("â†’", "→").replace("âš ï¸", "⚠️")
    text = text.replace("âœ…", "✅").replace("â³", "⏳")
    text = text.replace("Ã—", "×").replace("Ã ", "à").replace("Ã©", "é")
    text = text.replace("Ã¨", "è").replace("Ãª", "ê").replace("Ã¹", "ù")
    text = text.replace("Ã§", "ç").replace("Ã´", "ô").replace("Ã®", "î")
    text = text.replace("Ã¢", "â").replace("Ã»", "û").replace("Ã¯", "ï")
    text = text.replace("Â", "").replace("â€", "\"").replace("â€œ", "\"")
    text = text.replace("â€™", "'").replace("â€¢", "•")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    for level, size in [(1, 20), (2, 16), (3, 13), (4, 11.5)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(10 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run("Cahier des Charges V2")
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plateforme WhatsApp — Challenge Amazon FBA")
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(
        "Document de cadrage fonctionnel et opérationnel\nVersion client pour validation"
    )
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.italic = True

    box = doc.add_table(rows=3, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(2.0)
    box.columns[1].width = Inches(3.8)
    metadata = [
        ("Projet", "Challenge Amazon FBA"),
        ("Source", "CAHIER-DES-CHARGES-V2.md"),
        ("Généré le", datetime.now().strftime("%d/%m/%Y à %H:%M")),
    ]
    for row, (label, value) in zip(box.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], ACCENT_LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
    set_table_borders(box)
    doc.add_page_break()


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("Challenge Amazon FBA — Cahier des charges V2")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Document de travail client")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        tc = OxmlElement("w:shd")
        tc.set(qn("w:fill"), GRAY_LIGHT)
        p_pr = p._p.get_or_add_pPr()
        p_pr.append(tc)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = clean_inline(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
            if r_idx == 0:
                set_cell_shading(cell, ACCENT_LIGHT)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph("")


def parse_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(clean_inline(lines[i]))
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(level, 4)
            heading = clean_inline(stripped[level:].strip())
            doc.add_paragraph(heading, style=f"Heading {level}")
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(clean_inline(stripped[1:].strip()))
            run.italic = True
            run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            i += 1
            continue

        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?[\-\s|:]+\|?\s*$", lines[i + 1]):
            table_rows = [ [c.strip() for c in stripped.strip("|").split("|")] ]
            i += 2
            while i < len(lines):
                candidate = lines[i].strip()
                if "|" not in candidate or not candidate:
                    break
                table_rows.append([c.strip() for c in candidate.strip("|").split("|")])
                i += 1
            add_table(doc, table_rows)
            continue

        bullet_match = re.match(r"^([-*])\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="Normal")
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            run = p.add_run("• " + clean_inline(bullet_match.group(2)))
            run.font.size = Pt(10.5)
            i += 1
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.add_run(f"{numbered_match.group(1)}. {clean_inline(numbered_match.group(2))}")
            i += 1
            continue

        p = doc.add_paragraph(style="Normal")
        p.add_run(clean_inline(stripped))
        i += 1


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_header_footer(doc)
    parse_markdown(doc, text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
